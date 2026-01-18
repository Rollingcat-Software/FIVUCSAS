#!/usr/bin/env python3
"""
Generate Mermaid diagrams and create Word document with images
for FIVUCSAS ADD Document
"""

import re
import subprocess
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def extract_mermaid_diagrams(md_content):
    """Extract all Mermaid diagram blocks from markdown"""
    pattern = r'```mermaid\n(.*?)```'
    matches = re.findall(pattern, md_content, re.DOTALL)
    return matches


def generate_mermaid_image(mermaid_code, output_path, diagram_name):
    """Generate PNG image from Mermaid code using mmdc"""
    # Create a temporary .mmd file
    with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False, encoding='utf-8') as f:
        f.write(mermaid_code)
        mmd_file = f.name

    # Full path to mmdc on Windows
    mmdc_path = r'C:\Users\ahabg\AppData\Roaming\npm\mmdc.cmd'

    try:
        # Run mmdc to generate PNG
        result = subprocess.run(
            [mmdc_path, '-i', mmd_file, '-o', str(output_path), '-b', 'white', '-w', '1200'],
            capture_output=True,
            text=True,
            timeout=60,
            shell=True
        )

        if result.returncode != 0:
            print(f"  Warning: Failed to generate {diagram_name}: {result.stderr}")
            return False

        return output_path.exists()
    except subprocess.TimeoutExpired:
        print(f"  Warning: Timeout generating {diagram_name}")
        return False
    except Exception as e:
        print(f"  Warning: Error generating {diagram_name}: {e}")
        return False
    finally:
        # Clean up temp file
        Path(mmd_file).unlink(missing_ok=True)


def get_diagram_title(mermaid_code, index):
    """Try to determine a meaningful title for the diagram"""
    # Check for common diagram types
    if 'graph TB' in mermaid_code or 'graph LR' in mermaid_code:
        if 'Client Layer' in mermaid_code:
            return "System Architecture Overview"
        if 'Actors' in mermaid_code:
            return "Use Cases by Actor"
        if 'Git Submodules' in mermaid_code:
            return "Submodule Dependencies"
        if 'Repository' in mermaid_code:
            return "Repository Structure"
        if 'Docker Network' in mermaid_code:
            return "Docker Deployment"

    if 'sequenceDiagram' in mermaid_code:
        if 'Enroll' in mermaid_code:
            return "Face Enrollment Sequence"
        if 'Verify' in mermaid_code or 'Verification' in mermaid_code:
            return "Face Verification Sequence"
        if 'login' in mermaid_code.lower():
            return "Authentication Sequence"
        if 'Liveness' in mermaid_code or 'Challenge' in mermaid_code:
            return "Liveness Challenge Sequence"
        if 'Tenant' in mermaid_code:
            return "Multi-Tenant Request Flow"

    if 'classDiagram' in mermaid_code:
        if 'Tenant' in mermaid_code and 'User' in mermaid_code:
            return "Domain Model - Core Entities"
        if 'UseCase' in mermaid_code:
            return "Application Layer - Use Cases"

    if 'stateDiagram' in mermaid_code:
        if 'Pending' in mermaid_code and 'Active' in mermaid_code:
            return "User Lifecycle States"
        if 'FaceDetection' in mermaid_code:
            return "Verification Session States"
        if 'ChallengeGenerated' in mermaid_code:
            return "Biometric Puzzle States"

    if 'erDiagram' in mermaid_code:
        return "Entity Relationship Diagram"

    if 'flowchart' in mermaid_code:
        return "Microservices Communication"

    return f"Diagram {index + 1}"


def create_document():
    """Create and configure the Word document"""
    doc = Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Configure styles
    styles = doc.styles

    # Title style
    title_style = styles['Title']
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.name = 'Calibri'

    # Heading styles
    for level in range(1, 4):
        style = styles[f'Heading {level}']
        style.font.bold = True
        style.font.name = 'Calibri'
        if level == 1:
            style.font.size = Pt(16)
            style.paragraph_format.space_before = Pt(18)
        elif level == 2:
            style.font.size = Pt(14)
            style.paragraph_format.space_before = Pt(12)
        else:
            style.font.size = Pt(12)
            style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    # Normal text
    normal_style = styles['Normal']
    normal_style.font.size = Pt(11)
    normal_style.font.name = 'Calibri'
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal_style.paragraph_format.space_after = Pt(6)

    return doc


def parse_markdown_with_diagrams(md_content, diagram_images):
    """Parse markdown and replace mermaid blocks with image references"""
    lines = md_content.split('\n')
    elements = []
    current_table = None
    in_code_block = False
    code_content = []
    code_lang = None
    diagram_index = 0

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                if code_lang == 'mermaid':
                    # Insert diagram image reference
                    if diagram_index < len(diagram_images):
                        img_path, img_title = diagram_images[diagram_index]
                        if img_path and img_path.exists():
                            elements.append(('image', (img_path, img_title)))
                        diagram_index += 1
                else:
                    elements.append(('code', '\n'.join(code_content)))
                code_content = []
                in_code_block = False
                code_lang = None
            else:
                in_code_block = True
                # Check for language
                lang_match = re.match(r'^```(\w+)', line.strip())
                code_lang = lang_match.group(1) if lang_match else None
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # Skip empty lines
        if not line.strip():
            if current_table:
                elements.append(('table', current_table))
                current_table = None
            i += 1
            continue

        # Horizontal rules
        if line.strip() in ['---', '***', '___']:
            if current_table:
                elements.append(('table', current_table))
                current_table = None
            elements.append(('hr', None))
            i += 1
            continue

        # Headers
        if line.startswith('#'):
            if current_table:
                elements.append(('table', current_table))
                current_table = None

            match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if match:
                level = len(match.group(1))
                text = match.group(2).strip()
                elements.append((f'h{level}', text))
            i += 1
            continue

        # Tables
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]

            # Check if this is a separator line
            if all(re.match(r'^[-:]+$', c) for c in cells if c):
                i += 1
                continue

            if current_table is None:
                current_table = []
            current_table.append(cells)
            i += 1
            continue

        # List items
        if re.match(r'^\s*[-*+]\s+', line) or re.match(r'^\s*\d+\.\s+', line):
            if current_table:
                elements.append(('table', current_table))
                current_table = None

            if re.match(r'^\s*\d+\.\s+', line):
                match = re.match(r'^\s*\d+\.\s+(.+)$', line)
                if match:
                    elements.append(('ol', match.group(1)))
            else:
                match = re.match(r'^\s*[-*+]\s+(.+)$', line)
                if match:
                    elements.append(('ul', match.group(1)))
            i += 1
            continue

        # Regular paragraph
        if current_table:
            elements.append(('table', current_table))
            current_table = None

        para_lines = [line]
        while i + 1 < len(lines):
            next_line = lines[i + 1]
            if (not next_line.strip() or
                next_line.startswith('#') or
                next_line.strip().startswith('|') or
                next_line.strip().startswith('```') or
                re.match(r'^\s*[-*+]\s+', next_line) or
                re.match(r'^\s*\d+\.\s+', next_line)):
                break
            para_lines.append(next_line)
            i += 1

        elements.append(('p', ' '.join(para_lines)))
        i += 1

    if current_table:
        elements.append(('table', current_table))

    return elements


def clean_text(text):
    """Remove markdown formatting from text"""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text


def add_table(doc, table_data):
    """Add a table to the document"""
    if not table_data or not table_data[0]:
        return

    num_cols = len(table_data[0])
    table = doc.add_table(rows=len(table_data), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = clean_text(cell_text)

                if i == 0:
                    set_cell_shading(cell, 'D9E2F3')
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()


def add_image_with_caption(doc, img_path, caption):
    """Add an image with a centered caption"""
    # Add image
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()

    # Calculate width to fit page (max 15cm)
    run.add_picture(str(img_path), width=Cm(15))

    # Add caption
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run(f"Figure: {caption}")
    caption_run.italic = True
    caption_run.font.size = Pt(10)

    # Add space after
    doc.add_paragraph()


def convert_with_diagrams(md_path, docx_path, diagrams_dir):
    """Main conversion function with diagram generation"""
    print(f"Reading markdown from: {md_path}")

    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Extract and generate diagrams
    print("Extracting Mermaid diagrams...")
    mermaid_diagrams = extract_mermaid_diagrams(md_content)
    print(f"Found {len(mermaid_diagrams)} Mermaid diagrams")

    diagram_images = []
    for i, diagram_code in enumerate(mermaid_diagrams):
        title = get_diagram_title(diagram_code, i)
        img_name = f"diagram_{i+1:02d}_{title.lower().replace(' ', '_').replace('-', '_')[:30]}.png"
        img_path = diagrams_dir / img_name

        print(f"  Generating {i+1}/{len(mermaid_diagrams)}: {title}...")
        success = generate_mermaid_image(diagram_code, img_path, title)

        if success:
            diagram_images.append((img_path, title))
        else:
            diagram_images.append((None, title))

    successful = sum(1 for p, _ in diagram_images if p and p.exists())
    print(f"Successfully generated {successful}/{len(mermaid_diagrams)} diagrams")

    # Create document
    print("Creating Word document...")
    doc = create_document()

    print("Parsing markdown content...")
    elements = parse_markdown_with_diagrams(md_content, diagram_images)

    print(f"Processing {len(elements)} elements...")

    for elem_type, content in elements:
        if elem_type == 'h1':
            doc.add_heading(clean_text(content), level=1)
        elif elem_type == 'h2':
            doc.add_heading(clean_text(content), level=2)
        elif elem_type == 'h3':
            doc.add_heading(clean_text(content), level=3)
        elif elem_type == 'h4':
            p = doc.add_paragraph()
            run = p.add_run(clean_text(content))
            run.bold = True
            run.font.size = Pt(11)
        elif elem_type == 'h5' or elem_type == 'h6':
            p = doc.add_paragraph()
            run = p.add_run(clean_text(content))
            run.bold = True
            run.italic = True
        elif elem_type == 'p':
            p = doc.add_paragraph()
            p.add_run(clean_text(content))
        elif elem_type == 'ul':
            doc.add_paragraph(clean_text(content), style='List Bullet')
        elif elem_type == 'ol':
            doc.add_paragraph(clean_text(content), style='List Number')
        elif elem_type == 'table':
            add_table(doc, content)
        elif elem_type == 'code':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(content)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        elif elem_type == 'image':
            img_path, img_title = content
            add_image_with_caption(doc, img_path, img_title)
        elif elem_type == 'hr':
            doc.add_paragraph('_' * 80)

    print(f"Saving document to: {docx_path}")
    doc.save(docx_path)
    print("Conversion complete!")

    return docx_path


if __name__ == '__main__':
    base_dir = Path(__file__).parent.parent
    md_file = base_dir / 'docs' / 'ADD_FIVUCSAS.md'
    docx_file = base_dir / 'docs' / 'ADD_FIVUCSAS.docx'
    diagrams_dir = base_dir / 'docs' / 'ADD_diagrams'

    # Ensure diagrams directory exists
    diagrams_dir.mkdir(exist_ok=True)

    convert_with_diagrams(md_file, docx_file, diagrams_dir)
