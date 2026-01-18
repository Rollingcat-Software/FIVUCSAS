#!/usr/bin/env python3
"""
Markdown to Word Converter for FIVUCSAS ADD Document
Converts ADD_FIVUCSAS.md to a professionally formatted Word document
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def create_document():
    """Create and configure the Word document"""
    doc = Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Configure styles
    styles = doc.styles

    # Title style
    title_style = styles['Title']
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.name = 'Calibri'

    # Heading 1
    h1_style = styles['Heading 1']
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.name = 'Calibri'
    h1_style.paragraph_format.space_before = Pt(18)
    h1_style.paragraph_format.space_after = Pt(6)

    # Heading 2
    h2_style = styles['Heading 2']
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.name = 'Calibri'
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(6)

    # Heading 3
    h3_style = styles['Heading 3']
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.name = 'Calibri'
    h3_style.paragraph_format.space_before = Pt(10)
    h3_style.paragraph_format.space_after = Pt(4)

    # Normal text
    normal_style = styles['Normal']
    normal_style.font.size = Pt(11)
    normal_style.font.name = 'Calibri'
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal_style.paragraph_format.space_after = Pt(6)

    return doc


def parse_markdown(md_content):
    """Parse markdown content into structured elements"""
    lines = md_content.split('\n')
    elements = []
    current_table = None
    in_code_block = False
    code_content = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                elements.append(('code', '\n'.join(code_content)))
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # Skip empty lines
        if not line.strip():
            if current_table:
                elements.append(('table', current_table))
                current_table = None
            i += 1
            continue

        # Horizontal rules
        if line.strip() in ['---', '***', '___']:
            if current_table:
                elements.append(('table', current_table))
                current_table = None
            elements.append(('hr', None))
            i += 1
            continue

        # Headers
        if line.startswith('#'):
            if current_table:
                elements.append(('table', current_table))
                current_table = None

            match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if match:
                level = len(match.group(1))
                text = match.group(2).strip()
                elements.append((f'h{level}', text))
            i += 1
            continue

        # Tables
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]

            # Check if this is a separator line
            if all(re.match(r'^[-:]+$', c) for c in cells if c):
                i += 1
                continue

            if current_table is None:
                current_table = []
            current_table.append(cells)
            i += 1
            continue

        # List items
        if re.match(r'^\s*[-*+]\s+', line) or re.match(r'^\s*\d+\.\s+', line):
            if current_table:
                elements.append(('table', current_table))
                current_table = None

            # Determine list type and content
            if re.match(r'^\s*\d+\.\s+', line):
                match = re.match(r'^\s*\d+\.\s+(.+)$', line)
                if match:
                    elements.append(('ol', match.group(1)))
            else:
                match = re.match(r'^\s*[-*+]\s+(.+)$', line)
                if match:
                    elements.append(('ul', match.group(1)))
            i += 1
            continue

        # Regular paragraph
        if current_table:
            elements.append(('table', current_table))
            current_table = None

        # Collect multi-line paragraph
        para_lines = [line]
        while i + 1 < len(lines):
            next_line = lines[i + 1]
            if (not next_line.strip() or
                next_line.startswith('#') or
                next_line.strip().startswith('|') or
                next_line.strip().startswith('```') or
                re.match(r'^\s*[-*+]\s+', next_line) or
                re.match(r'^\s*\d+\.\s+', next_line)):
                break
            para_lines.append(next_line)
            i += 1

        elements.append(('p', ' '.join(para_lines)))
        i += 1

    # Handle any remaining table
    if current_table:
        elements.append(('table', current_table))

    return elements


def clean_text(text):
    """Remove markdown formatting from text"""
    # Remove bold
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    # Remove italic
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    # Remove code
    text = re.sub(r'`(.+?)`', r'\1', text)
    # Remove links but keep text
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text


def add_formatted_text(paragraph, text):
    """Add text with basic markdown formatting to a paragraph"""
    # Simple approach: just add cleaned text
    paragraph.add_run(clean_text(text))


def add_table(doc, table_data):
    """Add a table to the document"""
    if not table_data or not table_data[0]:
        return

    num_cols = len(table_data[0])
    table = doc.add_table(rows=len(table_data), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = clean_text(cell_text)

                # Header row styling
                if i == 0:
                    set_cell_shading(cell, 'D9E2F3')
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                # Cell text formatting
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(2)

    # Add space after table
    doc.add_paragraph()


def convert_md_to_docx(md_path, docx_path):
    """Main conversion function"""
    print(f"Reading markdown from: {md_path}")

    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    print("Creating Word document...")
    doc = create_document()

    print("Parsing markdown content...")
    elements = parse_markdown(md_content)

    print(f"Processing {len(elements)} elements...")

    for elem_type, content in elements:
        if elem_type == 'h1':
            doc.add_heading(clean_text(content), level=1)
        elif elem_type == 'h2':
            doc.add_heading(clean_text(content), level=2)
        elif elem_type == 'h3':
            doc.add_heading(clean_text(content), level=3)
        elif elem_type == 'h4':
            p = doc.add_paragraph()
            run = p.add_run(clean_text(content))
            run.bold = True
            run.font.size = Pt(11)
        elif elem_type == 'h5' or elem_type == 'h6':
            p = doc.add_paragraph()
            run = p.add_run(clean_text(content))
            run.bold = True
            run.italic = True
        elif elem_type == 'p':
            p = doc.add_paragraph()
            add_formatted_text(p, content)
        elif elem_type == 'ul':
            p = doc.add_paragraph(clean_text(content), style='List Bullet')
        elif elem_type == 'ol':
            p = doc.add_paragraph(clean_text(content), style='List Number')
        elif elem_type == 'table':
            add_table(doc, content)
        elif elem_type == 'code':
            # Add code as a formatted paragraph
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(content)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        elif elem_type == 'hr':
            # Add a line break
            doc.add_paragraph('_' * 80)

    print(f"Saving document to: {docx_path}")
    doc.save(docx_path)
    print("Conversion complete!")

    return docx_path


if __name__ == '__main__':
    # Paths
    base_dir = Path(__file__).parent.parent
    md_file = base_dir / 'docs' / 'ADD_FIVUCSAS.md'
    docx_file = base_dir / 'docs' / 'ADD_FIVUCSAS.docx'

    # Convert
    convert_md_to_docx(md_file, docx_file)
