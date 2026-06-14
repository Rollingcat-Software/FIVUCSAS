#!/usr/bin/env python3
"""Assemble the FIVUCSAS thesis Markdown chapters into the Software-Oriented .docx template.

Preserves the template's Guide-compliant format:
  - title/approval/abstract/ack/ToC/LoF/LoT front matter (roman page numbers)
  - body regenerated with ListParagraph + numId=2 auto-numbered headings (1 / 1.1 / 1.1.1)
  - each chapter on a new page (pageBreakBefore); page numbers restart at arabic 1 at Chapter 1
  - chapter-based Figure/Table caption SEQ fields so the List of Figures/Tables auto-populate
  - citations renumbered by first appearance; References + Appendices generated
Run:  python3 assemble.py
"""
import os, re, sys
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips, Emu
from docx.image.image import Image as DocxImage
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = os.environ.get("THESIS_ROOT") or os.path.dirname(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
T = ROOT + "/Thesis/build"
TEMPLATE = ROOT + "/Thesis/CSE4198_Thesis_Template_Software_Oriented_v2[1].docx"
OUT = os.environ.get("THESIS_OUT", ROOT + "/Thesis/FIVUCSAS_Thesis.docx")
CHAP_DIR = os.environ.get("THESIS_CHAP_DIR", T + "/chapters")
BIB = T + "/bibliography.md"

TITLE = "FACE AND IDENTITY VERIFICATION USING CLOUD-BASED SaaS MODELS"
AUTHORS = ["Ahmet Abdullah Gültekin", "Ayşe Gülsüm Eren", "Ayşenur Arıcı"]
SUPERVISOR = "Assoc. Prof. Dr. Mustafa Ağaoğlu"
YEAR = "2026"
SUBJECT = "CSE4298 Engineering Project, Marmara University, Computer Engineering"

# Page geometry (Guide: A4, 2.5 cm margins + 1 cm binding gutter). The usable text column
# is 11906 - 2*1418 - 567 = 8505 twips (15.0 cm); python-docx ignores the gutter when it
# sizes tables, so widths must be set explicitly.
A4_PGSZ = {"w:w": "11906", "w:h": "16838"}
BODY_PGMAR = {"w:top": "1418", "w:right": "1418", "w:bottom": "1418", "w:left": "1418",
              "w:header": "708", "w:footer": "709", "w:gutter": "567"}
TEXT_COL_TWIPS = 8505
MAX_FIG_W = Inches(5.9)    # full text-column width
MAX_FIG_H = Inches(8.85)   # one-page height cap, leaves room for the caption below
MIN_FIG_DPI = 220          # warn at build time when a raster lands below this

# ---------------------------------------------------------------- bibliography + figures
def load_bibliography():
    refs, figs = {}, {}
    section = None
    for line in open(BIB, encoding="utf-8"):
        s = line.strip()
        if s.startswith("## References"): section = "ref"; continue
        if s.startswith("## Figure"): section = "fig"; continue
        if section == "ref":
            m = re.match(r"- \*\*([A-Za-z0-9_\-]+)\*\* :: (.+)", s)
            if m: refs[m.group(1)] = m.group(2).strip()
        elif section == "fig":
            # | key | file | suggested caption | [optional display-width cap, e.g. 4.3in] |
            m = re.match(r"\|\s*([A-Za-z0-9_\-]+)\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|"
                         r"(?:\s*([0-9.]+)\s*in\s*\|)?", line)
            if m and m.group(1) not in ("key", "---", ":---"):
                width = Inches(float(m.group(4))) if m.group(4) else None
                figs[m.group(1)] = (m.group(2).strip(), m.group(3).strip(), width)
    return refs, figs

REFS, FIGS = load_bibliography()

# ---------------------------------------------------------------- markdown tokenizer
def tokenize(md):
    """Yield block tokens: ('h',level,title) ('p',text) ('bullet',text) ('num',text)
       ('fig',key,caption) ('table',caption,rows) ('eq',text) ('code',text)"""
    lines = md.split("\n")
    i, n = 0, len(lines)
    toks = []
    pending_cap = None  # caption captured from a [[TABLE: ...]] marker, applied to the NEXT table
    FIG_INLINE = r"\[\[FIG:\s*[A-Za-z0-9_\-]+\s*(?:\|\s*.*?)?\]\]"
    while i < n:
        st = lines[i].strip()
        if not st:
            i += 1; continue
        # code fence
        if st.startswith("```"):
            i += 1; buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i]); i += 1
            i += 1
            toks.append(("code", "\n".join(buf))); continue
        # heading
        m = re.match(r"^(#{1,4})\s+(.*)$", st)
        if m:
            level = len(m.group(1))
            title = re.sub(r"^\d+(\.\d+)*\.?\s*", "", m.group(2).strip()).strip()
            toks.append(("h", level, title)); i += 1; continue
        # markdown table block (caption = nearest preceding [[TABLE: ...]] marker)
        if st.startswith("|"):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(lines[i].strip()); i += 1
            toks.append(("table", pending_cap or "", rows)); pending_cap = None; continue
        # equation marker (the LaTeX line may be separated by blank lines)
        m = re.match(r"^\[\[EQ:\s*(.*?)\]\]$", st)
        if m:
            i += 1
            while i < n and not lines[i].strip():
                i += 1
            eq = ""
            if i < n and lines[i].strip() and not lines[i].strip().startswith("["):
                eq = lines[i].strip(); i += 1
            toks.append(("eq", m.group(1).strip(), eq)); continue
        # line that is ONLY figure marker(s) -> figure placement token(s)
        if re.match(r"^(?:" + FIG_INLINE + r"\s*)+$", st):
            for fm in re.finditer(r"\[\[FIG:\s*([A-Za-z0-9_\-]+)\s*(?:\|\s*(.*?))?\]\]", st):
                toks.append(("fig", fm.group(1), (fm.group(2) or "").strip()))
            i += 1; continue
        # standalone table caption marker -> remember for the next table
        m = re.match(r"^\[\[TABLE:\s*(.*?)\]\]$", st)
        if m:
            pending_cap = m.group(1).strip(); i += 1; continue
        # bullet / numbered (merge wrapped continuation lines like the paragraph branch:
        # a list item that wraps over several indented source lines is ONE item, not an
        # item followed by orphan justified paragraphs)
        def _consume_list_continuation(text, j):
            while j < n and lines[j].strip() and not re.match(
                    r"^(#{1,4}\s|\[\[FIG:|\[\[TABLE:|\[\[EQ:|[-*]\s|\d+[.)]\s|\||```)",
                    lines[j].strip()):
                text += " " + lines[j].strip(); j += 1
            return text, j
        m = re.match(r"^[-*]\s+(.*)$", st)
        if m:
            text, i = _consume_list_continuation(m.group(1).strip(), i + 1)
            toks.append(("bullet", text)); continue
        m = re.match(r"^\d+[.)]\s+(.*)$", st)
        if m:
            text, i = _consume_list_continuation(m.group(1).strip(), i + 1)
            toks.append(("num", text)); continue
        # paragraph (merge wrapped lines until blank/special); inline [[FIG]] stays in text
        buf = [st]; i += 1
        while i < n and lines[i].strip() and not re.match(
                r"^(#{1,4}\s|\[\[FIG:|\[\[TABLE:|\[\[EQ:|[-*]\s|\d+[.)]\s|\||```)", lines[i].strip()):
            buf.append(lines[i].strip()); i += 1
        ptext = " ".join(buf)
        # a trailing/inline [[TABLE: cap]] is the caption for the following table; pull it out
        tabs = re.findall(r"\[\[TABLE:\s*(.*?)\]\]", ptext)
        if tabs:
            pending_cap = tabs[-1].strip()
            ptext = re.sub(r"\s*\[\[TABLE:\s*.*?\]\]", "", ptext).strip()
        if ptext:
            toks.append(("p", ptext))
    return toks

def parse_table(rows):
    cells = []
    for r in rows:
        r = r.strip().strip("|")
        cols = [c.strip() for c in r.split("|")]
        cells.append(cols)
    # drop separator row (---)
    cells = [c for c in cells if not all(re.match(r"^:?-{2,}:?$", x.strip()) for x in c if x.strip()) or not c]
    cells = [c for c in cells if c and not all(re.match(r"^:?-{3,}:?$", x.strip() or "-") for x in c)]
    # simpler: remove rows that are all dashes
    out = []
    for c in rows:
        cc = [x.strip() for x in c.strip().strip("|").split("|")]
        if all(re.match(r"^:?-{2,}:?$", x) for x in cc if x != ""):
            continue
        out.append(cc)
    return out

# ---------------------------------------------------------------- citation numbering
def collect_citations(ordered_md):
    order, seen = [], set()
    for md in ordered_md:
        for m in re.finditer(r"\[CITE:\s*([A-Za-z0-9_,\-\s]+?)\]", md):
            for key in [k.strip() for k in m.group(1).split(",")]:
                if key and key not in seen:
                    seen.add(key); order.append(key)
    return {k: i + 1 for i, k in enumerate(order)}  # 1-based

def sub_citations(text, numbering):
    def repl(m):
        keys = [k.strip() for k in m.group(1).split(",") if k.strip()]
        nums = sorted({numbering.get(k) for k in keys if k in numbering})
        nums = [x for x in nums if x]
        return "[" + ",".join(str(x) for x in nums) + "]" if nums else ""
    return re.sub(r"\[CITE:\s*([A-Za-z0-9_,\-\s]+?)\]", repl, text)

# ---------------------------------------------------------------- oxml helpers
TNR = "Times New Roman"
def _set(el, tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), str(v))
    el.append(e); return e

def style_run_font(run, size_pt=12, bold=False, underline=False, italic=False, font=TNR, color=None):
    run.font.name = font; run.font.size = Pt(size_pt)
    run.font.bold = bold; run.font.underline = underline; run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.insert(0, rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), font)
    if color: run.font.color.rgb = color

def _para_mark_rpr(pPr, size_pt, bold, underline):
    """Paragraph-mark run properties. Word formats a numPr auto-number from these; without
    them the heading numbers fall back to docDefaults (Calibri 11 regular) while the heading
    text is TNR bold — the template's own headings carry the same pPr/rPr block."""
    rPr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), TNR)
    rPr.append(rf)
    if bold:
        rPr.append(OxmlElement("w:b"))
    for tag in ("w:sz", "w:szCs"):
        sz = OxmlElement(tag); sz.set(qn("w:val"), str(int(size_pt * 2))); rPr.append(sz)
    if underline:
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    pPr.append(rPr)  # rPr is the last formatting child of pPr

def add_heading(doc, title, ilvl, citation_numbering=None, page_break=False, numbered=True):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    _set(pPr, "w:pStyle", **{"w:val": "ListParagraph"})
    if numbered:
        numPr = OxmlElement("w:numPr")
        _set(numPr, "w:ilvl", **{"w:val": str(ilvl)})
        _set(numPr, "w:numId", **{"w:val": "2"})
        pPr.append(numPr)
    if page_break:
        _set(pPr, "w:pageBreakBefore")
    # Guide §Headings: chapter = UPPER/Bold/14pt/left-justified; L1 = Bold/12pt/indent 0.6cm;
    # L2 = Normal(not bold)/Underlined/12pt/indent 1.2cm. Hanging indent (425tw≈0.75cm) keeps the
    # auto-number at the required position with wrapped text aligned under the heading text.
    if ilvl == 0:
        _set(pPr, "w:spacing", **{"w:after": "120", "w:before": "240"})
        _set(pPr, "w:ind", **{"w:left": "425", "w:hanging": "425"})      # number at margin
    elif ilvl == 1:
        _set(pPr, "w:spacing", **{"w:after": "0", "w:line": "360", "w:lineRule": "auto"})
        _set(pPr, "w:ind", **{"w:left": "765", "w:hanging": "425"})      # number at 0.6 cm
    else:
        _set(pPr, "w:spacing", **{"w:after": "0", "w:line": "360", "w:lineRule": "auto"})
        _set(pPr, "w:ind", **{"w:left": "1105", "w:hanging": "425"})     # number at 1.2 cm
    _set(pPr, "w:jc", **{"w:val": "left"})
    _set(pPr, "w:outlineLvl", **{"w:val": str(ilvl)})
    _para_mark_rpr(pPr, size_pt=(14 if ilvl == 0 else 12),
                   bold=(ilvl < 2), underline=(ilvl >= 2))
    run = p.add_run(title.upper() if ilvl == 0 else title)
    style_run_font(run, size_pt=(14 if ilvl == 0 else 12),
                   bold=(ilvl < 2), underline=(ilvl >= 2))
    return p

def add_body(doc, text, citation_numbering, first_indent=True):
    text = sub_citations(text, citation_numbering)
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    _set(pPr, "w:jc", **{"w:val": "both"})
    _set(pPr, "w:spacing", **{"w:after": "120", "w:line": "360", "w:lineRule": "auto"})
    if first_indent:
        _set(pPr, "w:ind", **{"w:firstLine": "426"})
    add_inline_runs(p, text)
    return p

def add_inline_runs(p, text):
    # handle **bold**, *italic*, `code` — incl. `code` nested inside **bold**/*italic*
    def emit_with_code(seg_text, **fmt):
        for seg in re.split(r"(`[^`]+?`)", seg_text):
            if not seg: continue
            if seg.startswith("`") and seg.endswith("`"):
                r = p.add_run(seg[1:-1])
                style_run_font(r, font="Consolas", size_pt=10.5,
                               bold=fmt.get("bold", False), italic=fmt.get("italic", False))
            else:
                r = p.add_run(seg); style_run_font(r, **fmt)
    parts = re.split(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)", text)
    for part in parts:
        if not part: continue
        if part.startswith("**") and part.endswith("**"):
            emit_with_code(part[2:-2], bold=True)
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1]); style_run_font(r, font="Consolas", size_pt=10.5)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            emit_with_code(part[1:-1], italic=True)
        else:
            r = p.add_run(part); style_run_font(r)

def add_bullet(doc, text, citation_numbering, numbered=False, idx=1):
    text = sub_citations(text, citation_numbering)
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    _set(pPr, "w:jc", **{"w:val": "both"})
    _set(pPr, "w:spacing", **{"w:after": "60", "w:line": "360", "w:lineRule": "auto"})
    _set(pPr, "w:ind", **{"w:left": "720", "w:hanging": "360"})
    marker = (str(idx) + ".\t") if numbered else "•\t"
    r = p.add_run(marker); style_run_font(r)
    add_inline_runs(p, text)
    return p

def _run_color_auto(run):
    """Explicit automatic (black) run colour — the Caption style would otherwise paint the
    run theme-blue; the template's own example captions carry the same override."""
    c = run._element.get_or_add_rPr().get_or_add_color()
    c.set(qn("w:val"), "auto")

def add_caption(doc, label, chapter, is_first_in_chapter, text):
    """Caption paragraph: 'Figure C.' + SEQ(reset per chapter) + ' ' + text. Feeds TOC \\c.
    keepNext glues the caption to the next block: table captions precede their table and a
    figure paragraph (itself keepNext) precedes its caption, so neither can be orphaned."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    _set(pPr, "w:pStyle", **{"w:val": "Caption"})
    _set(pPr, "w:jc", **{"w:val": "center"})
    _set(pPr, "w:spacing", **{"w:before": "60", "w:after": "200"})
    p.paragraph_format.keep_with_next = True
    def run_text(s, bold=True):
        r = p.add_run(s); style_run_font(r, size_pt=10, bold=bold); _run_color_auto(r); return r
    run_text(label + " " + str(chapter) + ".")
    # SEQ field
    def fld(instr):
        r = p.add_run(); fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin"); r._r.append(fb)
        style_run_font(r, size_pt=10, bold=True); _run_color_auto(r)
        r2 = p.add_run(); it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr; r2._r.append(it)
        style_run_font(r2, size_pt=10, bold=True); _run_color_auto(r2)
        r3 = p.add_run(); fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate"); r3._r.append(fs)
        r4 = p.add_run("1"); style_run_font(r4, size_pt=10, bold=True); _run_color_auto(r4)
        r5 = p.add_run(); fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end"); r5._r.append(fe)
    fld(" SEQ %s \\* ARABIC %s " % (label, ("\\r 1" if is_first_in_chapter else "")))
    run_text("  " + text, bold=False)
    return p

def fit_figure_width(img, key, max_w=None):
    """Displayed width capped to the text column (or the catalog's per-figure cap) AND to
    one page of height (aspect kept), so tall diagrams shrink to fit instead of running
    off the page edge. Warns when the source raster falls below MIN_FIG_DPI at the
    resulting print size."""
    im = DocxImage.from_file(img)
    aspect = im.px_height / float(im.px_width)
    w = max_w if max_w and max_w < MAX_FIG_W else MAX_FIG_W
    if w * aspect > MAX_FIG_H:
        w = Emu(int(MAX_FIG_H / aspect))
    eff_dpi = im.px_width / w.inches
    if eff_dpi < MIN_FIG_DPI:
        sys.stderr.write("WARN: figure '%s' (%dx%d px at %.2f x %.2f in) is %d DPI (<%d) — "
                         "needs a higher-resolution export\n"
                         % (key, im.px_width, im.px_height, w.inches, w.inches * aspect,
                            eff_dpi, MIN_FIG_DPI))
    return w

def add_figure(doc, key, caption, chapter, is_first):
    path_cap = FIGS.get(key)
    if not path_cap:
        # unknown figure key -> placeholder
        add_caption(doc, "Figure", chapter, is_first, caption or ("[missing figure: %s]" % key)); return False
    rel, default_cap, fig_w = path_cap
    # vendored copy under build/figures first (self-contained build), then ROOT-relative,
    # then sibling checkouts (e.g. docs cloned next to the parent repo)
    candidates = [os.path.join(T, "figures", os.path.basename(rel)),
                  os.path.join(ROOT, rel), os.path.join(os.path.dirname(ROOT), rel)]
    img = next((c for c in candidates if os.path.exists(c)), candidates[1])
    p = doc.add_paragraph(); pPr = p._p.get_or_add_pPr()
    _set(pPr, "w:jc", **{"w:val": "center"})
    _set(pPr, "w:spacing", **{"w:before": "120", "w:after": "0"})
    p.paragraph_format.keep_with_next = True   # never page-break between figure and caption
    run = p.add_run()
    try:
        run.add_picture(img, width=fit_figure_width(img, key, fig_w))
    except Exception as e:
        run.add_text("[image error: %s]" % e)
    add_caption(doc, "Figure", chapter, is_first, caption or default_cap)
    return True

def size_table(tbl, ncols):
    """Fixed layout sized to the 8505-twip text column. python-docx derives its default
    table width from page-minus-margins and ignores the 1 cm binding gutter, which pushed
    every table ~1 cm past the right margin. Rows are made unsplittable and the header row
    repeats when a table runs over a page break."""
    tbl.autofit = False
    tblPr = tbl._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa"); tblW.set(qn("w:w"), str(TEXT_COL_TWIPS))
    base, rem = divmod(TEXT_COL_TWIPS, ncols)
    widths = [base + 1 if i < rem else base for i in range(ncols)]
    for col, w in zip(tbl.columns, widths):
        col.width = Twips(w)
    for ri, row in enumerate(tbl.rows):
        trPr = row._tr.get_or_add_trPr()
        _set(trPr, "w:cantSplit")          # a row never breaks across pages
        if ri == 0:
            _set(trPr, "w:tblHeader")      # header row repeats on every page
        for cell, w in zip(row.cells, widths):
            cell.width = Twips(w)

def add_table(doc, caption, rows, chapter, is_first, citation_numbering):
    data = parse_table(rows)
    if caption:
        add_caption(doc, "Table", chapter, is_first, sub_citations(caption, citation_numbering))
    if not data: return
    ncols = max(len(r) for r in data)
    tbl = doc.add_table(rows=0, cols=ncols)
    tbl.style = "Table Grid"
    for ri, r in enumerate(data):
        cells = tbl.add_row().cells
        for ci in range(ncols):
            txt = sub_citations(r[ci], citation_numbering) if ci < len(r) else ""
            cell = cells[ci]; cell.text = ""
            para = cell.paragraphs[0]
            add_inline_runs(para, txt)
            for rr in para.runs:
                style_run_font(rr, size_pt=10, bold=(ri == 0))
    size_table(tbl, ncols)
    # Word glues the paragraph that follows a table flush against its bottom border. A small
    # empty spacer paragraph after the table restores breathing room before the next text.
    spacer = doc.add_paragraph(); sp = spacer._p.get_or_add_pPr()
    _set(sp, "w:spacing", **{"w:after": "0", "w:before": "0", "w:line": "120", "w:lineRule": "exact"})
    return tbl

def add_code(doc, text):
    p = doc.add_paragraph(); pPr = p._p.get_or_add_pPr()
    _set(pPr, "w:spacing", **{"w:after": "120", "w:line": "240", "w:lineRule": "auto"})
    _set(pPr, "w:ind", **{"w:left": "360"})
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F2F2F2"); pPr.append(shd)
    for j, line in enumerate(text.split("\n")):
        if j: p.add_run().add_break()
        r = p.add_run(line); style_run_font(r, font="Consolas", size_pt=9)
    return p

# ---- OMML (Word native math) builders -------------------------------------
def _m(tag, parent=None):
    e = OxmlElement("m:" + tag)
    if parent is not None: parent.append(e)
    return e

def m_run(parent, text):
    r = _m("r", parent)
    rPr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), "Cambria Math"); rf.set(qn("w:hAnsi"), "Cambria Math")
    rPr.append(rf); r.append(rPr)
    t = _m("t", r); t.set(qn("xml:space"), "preserve"); t.text = text
    return r

def m_frac(parent, num_fn, den_fn):
    f = _m("f", parent)
    num_fn(_m("num", f)); den_fn(_m("den", f))
    return f

def m_ssub(parent, base, sub):
    s = _m("sSub", parent)
    m_run(_m("e", s), base); m_run(_m("sub", s), sub)
    return s

def _eq_ear(om):
    m_run(om, "EAR = ")
    m_frac(om, lambda n: m_run(n, "‖p₂ − p₆‖ + ‖p₃ − p₅‖"),
               lambda d: m_run(d, "2 ‖p₁ − p₄‖"))

def _eq_mar(om):
    m_run(om, "MAR = ")
    m_frac(om, lambda n: m_run(n, "‖lower_lip − upper_lip‖"),
               lambda d: m_run(d, "‖right_corner − left_corner‖"))

def _eq_cosine(om):
    m_ssub(om, "d", "cos")
    m_run(om, "(A, B) = 1 − ")
    m_frac(om, lambda n: m_run(n, "A · B"),
               lambda d: m_run(d, "‖A‖ ‖B‖"))

def _eq_quality(om):
    m_run(om, "Q = 0.4 · blur + 0.3 · lighting + 0.3 · face_size")

def _eq_acer(om):
    m_run(om, "ACER = ")
    m_frac(om, lambda n: m_run(n, "APCER + BPCER"),
               lambda d: m_run(d, "2"))

EQUATION_BUILDERS = {
    "Eye Aspect Ratio (EAR)": _eq_ear,
    "Mouth Aspect Ratio (MAR)": _eq_mar,
    "Cosine distance for face matching": _eq_cosine,
    "Composite image-quality score": _eq_quality,
    "ACER": _eq_acer,
}

def _latex_to_plain(eq):
    """Fallback: readable Unicode rendering of simple LaTeX when no OMML builder exists."""
    s = eq.strip().strip("$").strip()
    s = re.sub(r"\\(?:text|mathrm)\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\[dt]?frac\{([^}]*)\}\{([^}]*)\}", r"(\1) / (\2)", s)
    s = (s.replace(r"\lVert", "‖").replace(r"\rVert", "‖")
          .replace(r"\cdot", "·").replace(r"\,", " ").replace("\\_", "_")
          .replace("  ", " "))
    return s

def add_equation(doc, name, eq, chapter, counter):
    counter[0] += 1
    p = doc.add_paragraph(); pPr = p._p.get_or_add_pPr()
    _set(pPr, "w:jc", **{"w:val": "center"})
    _set(pPr, "w:spacing", **{"w:before": "60", "w:after": "60"})
    # right tab for the "(Equation x.y)" label at the right edge of the text column
    tabs = OxmlElement("w:tabs"); _set(tabs, "w:tab", **{"w:val": "right", "w:pos": "8500"}); pPr.append(tabs)
    builder = EQUATION_BUILDERS.get(name)
    if builder:
        # oMathPara marks display math — the proper OOXML form, survives more converters
        omp = _m("oMathPara"); om = _m("oMath", omp); builder(om); p._p.append(omp)
    else:
        sys.stderr.write("WARN: no OMML builder for equation '%s' — plain-text fallback\n" % name)
        r = p.add_run(_latex_to_plain(eq) if eq else name); style_run_font(r, italic=True)
    r2 = p.add_run("\t(Equation %d.%d)" % (chapter, counter[0])); style_run_font(r2)
    return p

def normalize_caption_style(doc):
    """The template's Caption style inherits the theme blue. Run-level overrides keep the
    literal caption runs black, but renderers that recalculate the SEQ fields (LibreOffice)
    style the recomputed number from the style itself — so fix the style too."""
    try:
        st = doc.styles["Caption"]
    except KeyError:
        return
    c = st.element.get_or_add_rPr().get_or_add_color()
    for a in ("w:themeColor", "w:themeTint", "w:themeShade"):
        if c.get(qn(a)) is not None:
            del c.attrib[qn(a)]
    c.set(qn("w:val"), "auto")

# ---------------------------------------------------------------- page geometry
def normalize_page_geometry(doc):
    """The template ships the title/approval sections as US Letter (12240x15840) while the
    rest is A4 — the Guide requires A4 throughout, and Letter is 2.5 cm shorter, which can
    overflow the title/copyright block. Force every section to A4; sections converted from
    Letter also take the body margin block (2.5 cm + 1 cm gutter)."""
    for sectPr in doc.element.body.iter(qn("w:sectPr")):
        pgSz = sectPr.find(qn("w:pgSz"))
        if pgSz is None:
            continue
        if (pgSz.get(qn("w:w")), pgSz.get(qn("w:h"))) == (A4_PGSZ["w:w"], A4_PGSZ["w:h"]):
            continue
        for k, v in A4_PGSZ.items():
            pgSz.set(qn(k), v)
        pgMar = sectPr.find(qn("w:pgMar"))
        if pgMar is not None:
            for k, v in BODY_PGMAR.items():
                pgMar.set(qn(k), v)

def drop_orphan_sign(doc):
    """The template pairs the supervisor AND the optional co-advisor each with a 'Sign'
    line on the Approval page; we blank the co-advisor lines, so his 'Sign' placeholder
    would survive as a floating line. Keep only the supervisor's."""
    signs = [p for p in doc.paragraphs if p.text.strip() == "Sign"]
    for p in signs[1:]:
        p._p.getparent().remove(p._p)

# ---------------------------------------------------------------- body section reset
def remove_body_after_frontmatter(doc):
    body = doc.element.body
    # locate first ListParagraph heading (INTRODUCTION) among body children
    children = list(body)
    start = None
    for idx, ch in enumerate(children):
        if ch.tag == qn("w:p"):
            txt = "".join(t.text or "" for t in ch.iter(qn("w:t")))
            ps = ch.find(qn("w:pPr"))
            sty = ps.find(qn("w:pStyle")) if ps is not None else None
            if sty is not None and sty.get(qn("w:val")) == "ListParagraph" and "INTRODUCTION" in txt.upper():
                start = idx; break
    assert start is not None, "could not find INTRODUCTION heading"
    trailing = children[-1] if children[-1].tag == qn("w:sectPr") else None
    # remove everything from start to end except trailing sectPr
    for ch in children[start:]:
        if ch is trailing: continue
        body.remove(ch)
    return trailing

def configure_body_section(trailing_sectPr):
    """Make the single body section arabic, restart at 1, footer rId17 (footer with PAGE field)."""
    if trailing_sectPr is None: return
    # remove existing pgNumType / footerReference(default)
    for tag in ("w:pgNumType",):
        for e in trailing_sectPr.findall(qn(tag)):
            trailing_sectPr.remove(e)
    # set footer default -> rId17 (footer3 has PAGE field)
    for fr in trailing_sectPr.findall(qn("w:footerReference")):
        if fr.get(qn("w:type")) == "default":
            trailing_sectPr.remove(fr)
    fr = OxmlElement("w:footerReference"); fr.set(qn("w:type"), "default"); fr.set(qn("r:id"), "rId17")
    trailing_sectPr.insert(0, fr)
    png = OxmlElement("w:pgNumType"); png.set(qn("w:start"), "1")
    # insert pgNumType after pgMar
    pgmar = trailing_sectPr.find(qn("w:pgMar"))
    if pgmar is not None: pgmar.addnext(png)
    else: trailing_sectPr.append(png)

# ---------------------------------------------------------------- front matter text replace
def replace_frontmatter(doc, abstract, acknowledgements):
    for p in doc.paragraphs:
        txt = p.text.replace(" ", " ")
        def setp(newtext, **fmt):
            for r in list(p.runs):
                r._r.getparent().remove(r._r)
            r = p.add_run(newtext); style_run_font(r, **fmt)
        if "Copyright" in txt:
            setp("Copyright © " + ", ".join(AUTHORS) + ", " + YEAR + ". All rights reserved.")
        elif "InsERT TITLE OF YOUR PROJECT" in txt:
            setp(TITLE, bold=True, size_pt=14)
        elif "Insert second line of title" in txt:
            setp("")  # no second line
        elif "Group member" in txt:
            # fill three author lines in order via a counter on the function attribute
            replace_frontmatter._ai = getattr(replace_frontmatter, "_ai", 0)
            idx = replace_frontmatter._ai % 3
            setp(AUTHORS[idx] if idx < len(AUTHORS) else "")
            replace_frontmatter._ai += 1
        elif "Insert supervisor's title" in txt:
            setp(SUPERVISOR)
        elif "Insert co-advisor" in txt or "Co-advised by" in txt:
            setp("")
        elif "CSE4197" in txt:
            # The template boilerplate reads "CSE4197 / CSE4198"; this team's course codes
            # are CSE4297 (fall PSD term) / CSE4298 (thesis term). The digits sit in their
            # own runs, so flip just the "1" runs and keep the template formatting intact.
            runs = p.runs
            for _i in range(1, len(runs)):
                if runs[_i].text == "1" and runs[_i - 1].text.endswith("CSE4"):
                    runs[_i].text = "2"
            if "CSE4197" in p.text:  # fallback if the template's run layout ever changes
                for r in p.runs:
                    r.text = r.text.replace("4197", "4297").replace("4198", "4298")
        elif "Year of Graduation" in txt:
            setp(YEAR)
        elif txt.strip().startswith("The abstract part is a brief summary"):
            setp(abstract)
            _ppr = p._p.get_or_add_pPr()
            for _ind in _ppr.findall(qn("w:ind")): _ppr.remove(_ind)   # drop template's 708 twips
            _set(_ppr, "w:ind", **{"w:firstLine": "426"})              # 0.75cm per Guide
        elif txt.strip().startswith("This part is optional"):
            setp(acknowledgements)
            _ppr = p._p.get_or_add_pPr()
            for _ind in _ppr.findall(qn("w:ind")): _ppr.remove(_ind)
            _set(_ppr, "w:ind", **{"w:firstLine": "426"})

# ---------------------------------------------------------------- main build
FIG_RE = r"\[\[FIG:\s*([A-Za-z0-9_\-]+)\s*(?:\|\s*(.*?))?\]\]"

def number_figures(toks):
    """Per-chapter figure number by FIRST appearance (inline-in-text or standalone)."""
    fignum = {}; c = 0
    for t in toks:
        if t[0] in ("p", "bullet", "num"):
            for m in re.finditer(FIG_RE, t[1]):
                if m.group(1) not in fignum:
                    c += 1; fignum[m.group(1)] = c
        elif t[0] == "fig":
            if t[1] not in fignum:
                c += 1; fignum[t[1]] = c
    return fignum

def render_chapter(doc, chapter_no, toks, citation_numbering):
    tbl_count = [0]; eq_count = [0]
    fignum = number_figures(toks)
    embedded = set(); first_fig = [True]; first_tbl = [True]

    def fig_ref(text):
        return re.sub(FIG_RE, lambda m: "Figure %d.%d" % (chapter_no, fignum.get(m.group(1), 0)), text)

    def embed(key, cap):
        if key in embedded:
            return
        embedded.add(key)
        add_figure(doc, key, cap, chapter_no, first_fig[0])
        first_fig[0] = False

    def text_token(render_fn, text):
        keys = [(m.group(1), (m.group(2) or "").strip()) for m in re.finditer(FIG_RE, text)]
        render_fn(fig_ref(text))
        for k, cap in keys:
            embed(k, cap)

    num_idx = 0   # position within a run of consecutive ordered-list items
    for t in toks:
        num_idx = num_idx + 1 if t[0] == "num" else 0
        if t[0] == "h":
            ilvl = t[1] - 1
            add_heading(doc, t[2], ilvl, page_break=(ilvl == 0))
        elif t[0] == "p":
            text_token(lambda s: add_body(doc, s, citation_numbering, first_indent=True), t[1])
        elif t[0] == "bullet":
            text_token(lambda s: add_bullet(doc, s, citation_numbering), t[1])
        elif t[0] == "num":
            text_token(lambda s, i=num_idx: add_bullet(doc, s, citation_numbering, numbered=True, idx=i), t[1])
        elif t[0] == "fig":
            embed(t[1], t[2])
        elif t[0] == "table":
            tbl_count[0] += 1
            add_table(doc, t[1], t[2], chapter_no, first_tbl[0], citation_numbering)
            first_tbl[0] = False
        elif t[0] == "eq":
            add_equation(doc, t[1], t[2], chapter_no, eq_count)
        elif t[0] == "code":
            add_code(doc, t[1])

def add_references(doc, citation_numbering):
    add_heading(doc, "REFERENCES", 0, page_break=True, numbered=False)
    inv = sorted(citation_numbering.items(), key=lambda kv: kv[1])
    for key, num in inv:
        ref = REFS.get(key, "[MISSING REFERENCE: %s]" % key)
        p = doc.add_paragraph(); pPr = p._p.get_or_add_pPr()
        _set(pPr, "w:jc", **{"w:val": "both"})
        _set(pPr, "w:spacing", **{"w:after": "120", "w:line": "276", "w:lineRule": "auto"})
        _set(pPr, "w:ind", **{"w:left": "567", "w:hanging": "567"})
        r = p.add_run("[%d] " % num); style_run_font(r, bold=True)
        # entries carry the Guide-mandated italics as *...* (journal/conference/book/web titles)
        add_inline_runs(p, ref)

def main():
    doc = Document(TEMPLATE)
    # read chapters
    def read(name):
        path = os.path.join(CHAP_DIR, name)
        return open(path, encoding="utf-8").read() if os.path.exists(path) else ""
    fm = read("00_frontmatter.md")
    chap_md = {i: read("ch%d.md" % i) for i in range(1, 8)}
    appendix_md = read("appendices.md")

    # extract abstract + acknowledgements from frontmatter
    abstract, ack = "", ""
    fm_secs = re.split(r"^#\s+", fm, flags=re.M)
    for sec in fm_secs:
        if sec.upper().startswith("ABSTRACT"):
            abstract = re.sub(r"^ABSTRACT\s*", "", sec, flags=re.I).strip()
        elif sec.upper().startswith("ACKNOWLEDG"):
            ack = re.sub(r"^ACKNOWLEDGE\w*\s*", "", sec, flags=re.I).strip()
    abstract = " ".join(abstract.split())
    ack = " ".join(ack.split())

    # citation numbering across body in order
    ordered = [chap_md[i] for i in range(1, 8)] + [appendix_md]
    citation_numbering = collect_citations(ordered)

    normalize_page_geometry(doc)
    normalize_caption_style(doc)
    replace_frontmatter(doc, abstract or "(abstract pending)", ack or "")
    drop_orphan_sign(doc)
    trailing = remove_body_after_frontmatter(doc)
    for i in range(1, 8):
        if chap_md[i].strip():
            render_chapter(doc, i, tokenize(chap_md[i]), citation_numbering)
    add_references(doc, citation_numbering)
    if appendix_md.strip():
        render_appendices(doc, appendix_md, citation_numbering)
    configure_body_section(trailing)
    # ask Word to refresh all fields (ToC / List of Figures / List of Tables / SEQ / PAGEREF) on open
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        uf = OxmlElement("w:updateFields"); uf.set(qn("w:val"), "true"); settings.append(uf)
    # document metadata: the template carries its own creator's name — stamp ours
    cp = doc.core_properties
    cp.author = "; ".join(AUTHORS)
    cp.last_modified_by = AUTHORS[0]
    cp.title = TITLE
    cp.subject = SUBJECT
    cp.modified = datetime.now(timezone.utc)
    doc.save(OUT)
    # stats
    wc_abs = len(abstract.split())
    print("SAVED:", OUT)
    print("abstract words:", wc_abs, "(target 150-300)")
    print("citations:", len(citation_numbering))
    missing = [k for k in citation_numbering if k not in REFS]
    print("missing refs:", missing if missing else "none")

def render_appendices(doc, md, citation_numbering):
    letters = iter("ABCDEFGHIJ")
    cur_letter = ["A"]; first_tbl = [True]
    num_idx = 0
    for t in tokenize(md):
        num_idx = num_idx + 1 if t[0] == "num" else 0
        if t[0] == "h" and t[1] == 1:
            cur_letter[0] = next(letters); first_tbl[0] = True
            add_heading(doc, t[2], 0, numbered=False, page_break=True)  # APPENDIX X — Title (new page)
            # Guide: "a cover page must precede each appendix" — title stands alone, body on the next page
            _cp = doc.add_paragraph(); _cr = _cp.add_run()
            _cbr = OxmlElement("w:br"); _cbr.set(qn("w:type"), "page"); _cr._r.append(_cbr)
        elif t[0] == "h":
            add_heading(doc, t[2], t[1] - 1, numbered=False)
        elif t[0] == "p":
            add_body(doc, t[1], citation_numbering)
        elif t[0] == "bullet":
            add_bullet(doc, t[1], citation_numbering)
        elif t[0] == "num":
            add_bullet(doc, t[1], citation_numbering, numbered=True, idx=num_idx)
        elif t[0] == "table":
            add_table(doc, t[1], t[2], cur_letter[0], first_tbl[0], citation_numbering)
            first_tbl[0] = False
        elif t[0] == "code":
            add_code(doc, t[1])

if __name__ == "__main__":
    main()
