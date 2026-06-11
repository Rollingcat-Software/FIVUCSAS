#!/usr/bin/env python3
import re, zipfile, sys, os
from docx import Document
DOCX = os.environ.get("THESIS_OUT") or os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "FIVUCSAS_Thesis.docx")
z = zipfile.ZipFile(DOCX)
doc = z.read("word/document.xml").decode("utf-8")
media = [n for n in z.namelist() if n.startswith("word/media/")]
paras = re.findall(r"<w:p\b.*?</w:p>", doc, re.S)
def text(p):
    return "".join(re.findall(r"<w:t[ >].*?</w:t>", p, re.S)).replace('<w:t xml:space="preserve">', "").replace("<w:t>", "").replace("</w:t>", "")
print("=== Chapter headings (numbered ilvl0) in order ===")
for p in paras:
    if re.search(r'pStyle w:val="ListParagraph"', p) and re.search(r'<w:ilvl w:val="0"', p) and re.search(r"<w:numId", p):
        print("   ", text(p)[:62])
print("=== unnumbered ilvl0 (REFERENCES/APPENDICES) ===")
for p in paras:
    if re.search(r'pStyle w:val="ListParagraph"', p) and re.search(r'<w:outlineLvl w:val="0"', p) and not re.search(r"<w:numId", p):
        print("   ", text(p)[:40])
print("\n=== counts ===")
print("  paragraphs:", len(paras))
print("  H1 numbered:", sum(1 for p in paras if re.search(r"ListParagraph", p) and re.search(r'<w:ilvl w:val="0"', p) and re.search(r"<w:numId", p)))
print("  H2:", sum(1 for p in paras if re.search(r'<w:ilvl w:val="1"', p)))
print("  H3:", sum(1 for p in paras if re.search(r'<w:ilvl w:val="2"', p)))
print("  tables:", doc.count("<w:tbl>"))
print("  SEQ Figure fields:", doc.count("SEQ Figure"))
print("  SEQ Table fields:", doc.count("SEQ Table"))
print("  equations:", doc.count("(Equation"))
print("  OMML math objects:", doc.count("<m:oMath>"))
print("  images embedded:", len(media))
print("\n=== leftover markers (must be 0) ===")
for m in ["[[FIG", "[[TABLE", "[[EQ", "[CITE:", "## ", "**", "```", "$$", "\\frac", "\\lVert", "`"]:
    print("  %-9r %d" % (m, doc.count(m)))
print("\n=== References ===")
refs = [text(p) for p in paras if re.match(r"^\[\d+\] ", text(p))]
print("  entries:", len(refs))
if refs: print("  first:", refs[0][:64]); print("  last :", refs[-1][:64])
d = Document(DOCX)
print("\n  python-docx reopen OK: paragraphs=%d tables=%d" % (len(d.paragraphs), len(d.tables)))
